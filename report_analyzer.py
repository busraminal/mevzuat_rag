# -*- coding: utf-8 -*-
# (1) Python 3.8+ için ileriye dönük type annotation'ları kolaylaştırır.
from __future__ import annotations

"""
CSV rapor analizörü — sayfa sayfa fark/hata özeti + HTML/MD/Excel (+opsiyonel DOCX)

Amaç:
- pdf_gap_checker.py'nin ürettiği CSV'yi okuyup son kullanıcı için anlaşılır bir özet çıkarmak.
- Her sorun satırı için:
    * KOD (explain_code): E101/E102/... gibi kısa ve sınıflandırılabilir bir kod
    * TESPİT (explain_text): insanın okuyacağı yalın açıklama
    * ÖNERİ (action_hint): ne yapılmalı? (düzeltme önerisi)
- Snippet'lardan madde/bent numarası yakalamak (madde_pdf / madde_law)
- diff_note içindeki liste benzeri değerleri ayrıştırmak (pdf_list / law_list)
- Sadece gerçekten "uyuşmayan" çiftleri görmek (mismatch_pairs)
- (yeni) Kodun adı/açıklaması (ecode_name/ecode_desc), satır sinyalleri (detected_kinds) ve
  kısa raporda kod dağılımı özeti

Kullanım:
    python report_analyzer.py --csv .\reports\report_YYYYMMDD_HHMM.csv --out .\reports [--law_version "RG 01/01/2024"]
"""

# (2) Standart kütüphaneler
import unicodedata  # TR aksanlarını ASCII'ye indirgemek için (yüzde->yuzde vb.)
import argparse
import os
import re
from datetime import datetime
from ast import literal_eval
from pathlib import Path
from typing import Optional, Tuple, List, Dict, Any

# (3) 3. parti
import pandas as pd


# =============================================================================
# SABİTLER ve REGEX'LER
# =============================================================================

# (4) Hangi status'ları "problem" sayalım?
# - "ambiguous"ı da problem sayıyoruz (insan bakışı gerekebilir).
# - "weak_aligned" isteğe göre dahil edilebilir; şimdilik değil.
PROBLEM_STATUSES = {"conflict", "missing", "low-similarity", "ambiguous"}
OK_STATUSES = {"aligned"}

# (5) 1, 1.2, 1.2.3.4 gibi "madde/alt madde" kalıbı — snippet'tan madde bulmak için.
_MADDE_PAT = re.compile(r"(?<!\d)(\d+(?:\.\d+){0,4})(?!\d)")

# (6) diff / snippet içinden sayı türlerini anlamaya yarayacak basit ipuçları:
_MONEY_PAT    = re.compile(r"(?:₺|TL|TRY)|\d{1,3}(?:[.,]\d{3})+(?:[.,]\d+)?")
_PERCENT_PAT  = re.compile(r"%\s*\d+(?:[.,]\d+)?")
_DATE_PAT     = re.compile(r"\b\d{1,2}[./]\d{1,2}[./]\d{2,4}\b")
_DURATION_PAT = re.compile(r"(\d+(?:[.,]\d+)?)\s*(gün|hafta|ay|yıl|saat|dk|dakika)\b", re.I)

# (6.1) Negasyon/istisna ve yön (izin/yasak) için basit sözlükler — heuristik amaçlı
NEGATION_TERMS = {
    "hariç", "istisna", "değildir", "olmadığı takdirde", "ancak", "yasaktır",
    "izin verilmeyecektir", "yasaklanır", "mümkün değildir", "yasak", "yapılamaz"
}
ALLOW_TERMS = {"izin verilir", "uygulanır", "mümkündür", "caizdir", "yapılabilir"}
FORBID_TERMS = {"yasaktır", "yasaklanır", "uygulanmaz", "yapılamaz", "men edilir"}

# (6.2) Kod sözlüğü — kısa ad/özet açıklama (MD/Excel’e ekliyoruz)
E_CODE_DICT: Dict[str, Dict[str, str]] = {
    "E101": {"ad": "Tutar farkı",                 "aciklama": "Parasal değerler uyuşmuyor."},
    "E102": {"ad": "Yüzde farkı",                 "aciklama": "Yüzde/oran değerleri farklı."},
    "E103": {"ad": "Süre farkı",                  "aciklama": "Aynı birimde süreler farklı."},
    "E104": {"ad": "Tarih farkı",                 "aciklama": "Tarihler uyuşmuyor."},
    "E105": {"ad": "Madde/bent atfı farklı",      "aciklama": "Madde/bent eşleşmesi farklı."},
    "E106": {"ad": "Kural yönü çelişkisi",        "aciklama": "izin/uygulanır ↔ yasak/uygulanmaz."},
    "E108": {"ad": "Düşük benzerlik",             "aciklama": "Anlamsal eşleşme zayıf."},
    "E109": {"ad": "Eksik içerik",                "aciklama": "Mevzuatta var, PDF'de yok."},
    "E110": {"ad": "Belirsiz eşleşme",            "aciklama": "Adaylar yakın."},
    "E111": {"ad": "Kısmi uyum",                  "aciklama": "İfadeler tam örtüşmüyor."},
    "E112": {"ad": "Versiyon/yürürlük sorunu",    "aciklama": "Yanlış mevzuat versiyonu."},
    "E113": {"ad": "OCR/format hatası",           "aciklama": "PDF okuma/biçim hatası."},
    "E000": {"ad": "Genel/Belirsiz",              "aciklama": "Tip belirlenemedi."},
}


# =============================================================================
# KOLON ADI KEŞFİ (esnek)
# =============================================================================

def smart_col(df: pd.DataFrame, name_cands: List[str]) -> str:
    """
    CSV'deki kolon adları farklı yazılmış olabilir (case / boşluk / tire / vs.)
    Bu yardımcı, aday isimlerden ilk eşleşeni döndürür; birebir eşleşme yoksa
    alfasayısal karakterlere indirgenmiş hâliyle dener.
    """
    # (7) case-insensitive doğrudan eşleştirme
    norm = {c.lower(): c for c in df.columns}
    for cand in name_cands:
        lc = cand.lower()
        if lc in norm:
            return norm[lc]

    # (8) "özel karakterleri at" yaklaşımıyla ikinci deneme
    def normkey(s: str) -> str:
        return re.sub(r"\W+", "", s.lower())

    target = normkey(name_cands[0])
    for c in df.columns:
        if normkey(c) == target:
            return c

    # (9) Artık bulunamadıysa hata
    raise KeyError(f"Kolon bulunamadı: {name_cands}")


def smart_col_opt(df: pd.DataFrame, name_cands: List[str]) -> Optional[str]:
    """smart_col'un opsiyonel versiyonu — yoksa None döndürür."""
    try:
        return smart_col(df, name_cands)
    except KeyError:
        return None


# =============================================================================
# METİN TEMİZLİK/YARDIMCI ARAÇLAR
# =============================================================================

def trim(s, n: int = 220) -> str:
    """
    Hücre içeriğini kısaltır ve whitespace normalize eder.
    Not: MD/HTML/Excel gibi farklı çıktılarda okunabilirliği artırır.
    """
    if not isinstance(s, str):
        s = "" if pd.isna(s) else str(s)
    s = " ".join(s.split())
    return s if len(s) <= n else s[:n] + " ..."


def md_escape(s: str) -> str:
    """Markdown tablolarda '|' karakteri sorun çıkarır, kaçırıyoruz."""
    return s.replace("|", r"\|")


def extract_madde(text: Optional[str]) -> Optional[str]:
    """Snippet içinden 1.2.3 gibi bir madde deseni yakalar; bulunamazsa None."""
    if not isinstance(text, str):
        return None
    m = _MADDE_PAT.search(text)
    return m.group(1) if m else None


# =============================================================================
# diff_note AYRIŞTIRMA ( '... [a,b,c] vs [x,y,z]' → iki liste )
# =============================================================================

def parse_diff_lists(diff_note: Optional[str]) -> Tuple[Optional[list], Optional[list]]:
    """
    diff_note alanında sık görülen kalıp:
      'tutar farkı: ['10.000 TL'] vs ['12.000 TL']'
    gibi. Buradan iki Python listesi (pdf_list, law_list) üretmeye çalışır.
    Robust olması için literal_eval + fallback string parçalama yapar.
    """
    if not isinstance(diff_note, str):
        return (None, None)

    # (10) 'vs' iki tarafı birbirinden ayırmak için kritik
    parts = diff_note.split("vs")
    if len(parts) != 2:
        return (None, None)

    def _find_list(s: str) -> Optional[list]:
        # Köşeli parantez içini yakala
        m = re.search(r"\[(.*?)\]", s, flags=re.S)
        if not m:
            return None

        # OCR gibi hatalı durumlar için minik düzeltmeler (OOTL → TL)
        raw = "[" + m.group(1) + "]"
        safe = raw.replace("OOTL", "TL")

        # Önce literal_eval ile deneriz (['a','b'] gibi clean listelerde başarılı)
        try:
            return literal_eval(safe)
        except Exception:
            # Olmazsa kaba parçalama
            try:
                L = [x.strip() for x in m.group(1).split(",")]
                return [x.strip("'").strip('"') for x in L]
            except Exception:
                return None

    return (_find_list(parts[0]), _find_list(parts[1]))


def mismatch_pairs(pdf_list: Optional[list], law_list: Optional[list]) -> List[Tuple[str, str]]:
    """
    İki listeden, yalnızca iki tarafta da dolu olup birbirinden FARKLI olan çiftleri döndür.
    (Boş ↔ dolu gibi tek taraflı durumları bu raporda göstermeyelim — noise azaltır.)
    """
    a = pdf_list or []
    b = law_list or []
    n = max(len(a), len(b))
    out: List[Tuple[str, str]] = []
    for i in range(n):
        va = "" if i >= len(a) else str(a[i]).strip()
        vb = "" if i >= len(b) else str(b[i]).strip()
        if not va or not vb:
            continue
        if va != vb:
            out.append((va, vb))
    return out


# =============================================================================
# TÜR/İPUCU TESPİTİ (money/percent/date/duration/negation/polarity)
# =============================================================================

def _count_hits(text: str, vocab: set) -> int:
    """Basit 'içerir mi' sayacı (heuristik)."""
    t = (text or "").lower()
    return sum(1 for w in vocab if w in t)

def detect_polarity_conflict(pdf_text: str, law_text: str) -> bool:
    """
    Heuristik: PDF 'izin' yönünde, mevzuat 'yasak' yönünde (veya tersi) ise True.
    Çok katı değil; basit anahtar kelime sayımıyla karar veriyoruz.
    """
    pdf_allow = _count_hits(pdf_text, ALLOW_TERMS)
    pdf_forbid = _count_hits(pdf_text, FORBID_TERMS)
    law_allow = _count_hits(law_text, ALLOW_TERMS)
    law_forbid = _count_hits(law_text, FORBID_TERMS)

    # PDF izin baskın & mevzuat yasak baskın → çelişki
    if pdf_allow > pdf_forbid and law_forbid > law_allow:
        return True
    # PDF yasak baskın & mevzuat izin baskın → çelişki
    if pdf_forbid > pdf_allow and law_allow > law_forbid:
        return True
    return False


def _kind_from_text(s: str) -> set:
    """Verilen metindeki sinyallere göre tür etiketi set'i döndürür."""
    kinds = set()
    if not isinstance(s, str) or not s.strip():
        return kinds
    if _MONEY_PAT.search(s):    kinds.add("money")
    if _PERCENT_PAT.search(s):  kinds.add("percent")
    if _DATE_PAT.search(s):     kinds.add("date")
    if _DURATION_PAT.search(s): kinds.add("duration")
    # Negasyon/istisna sinyali (salt bilgi; tek başına polarity demeyelim)
    if any(word in s.lower() for word in NEGATION_TERMS):
        kinds.add("has_negation")
    return kinds


def _kind_from_row(row: pd.Series) -> set:
    """
    Tek satır (problem kaydı) içindeki PDF ve Mevzuat snippet'larını ve diff_note'u
    inceleyip türleri birleştirir.
    """
    kinds = set()
    pdf_t = row.get("pdf_snippet", "") or ""
    law_t = row.get("law_snippet", "") or ""
    kinds |= _kind_from_text(pdf_t)
    kinds |= _kind_from_text(law_t)

    dn = str(row.get("diff_note", "")).lower()
    if "tutar farkı" in dn:  kinds.add("money")
    if "yüzde farkı" in dn:  kinds.add("percent")
    if "tarih farkı" in dn:  kinds.add("date")
    if "süre farkı"  in dn:  kinds.add("duration")

    # (ek) yön çelişkisi heuristiği
    if detect_polarity_conflict(pdf_t, law_t):
        kinds.add("polarity")
    return kinds


# =============================================================================
# SON KULLANICI: KOD + TESPİT + ÖNERİ üretimi
# =============================================================================

def make_explanation(row: pd.Series) -> tuple[str, str, str]:
    """
    Her problem satırı için 3'lü döndürür:
      (explain_code, explain_text, action_hint)
    - explain_code: E101/E102/... (raporlamada filtreleme/analitik için)
    - explain_text: son kullanıcıya "ne sorun var?"
    - action_hint : "ne yapmalı?"
    """
    status = str(row.get("status", "")).lower()
    dn = str(row.get("diff_note", "")).lower()

    # diff listelerini ve örnek çiftleri hazırla
    pdf_list, law_list = row.get("pdf_list"), row.get("law_list")
    pairs = mismatch_pairs(pdf_list, law_list)
    sample = "; ".join([f"{x} ↔ {y}" for x, y in pairs[:3]]) if pairs else ""

    # snippet'lardan yakalanmış madde numaraları
    madde_pdf = row.get("madde_pdf") or "-"
    madde_law = row.get("madde_law") or "-"

    # tür sinyalleri
    kinds = _kind_from_row(row)

    # ---- 1) Net sayısal/kategorik farklar
    if "tutar farkı" in dn or "money" in kinds:
        return ("E101",
                "Parasal değerler uyuşmuyor" + (f" (örn: {sample})" if sample else ""),
                "PDF’deki tutar(ları) mevzuata göre güncelleyin; para birimi/ayraçları normalize edin (örn. 10.000,00 ₺ → 10.000 TL).")

    if "yüzde farkı" in dn or "percent" in kinds:
        return ("E102",
                "Yüzde oranları uyuşmuyor" + (f" (örn: {sample})" if sample else ""),
                "PDF’deki yüzde(leri) mevzuattaki oranlara hizalayın; yuvarlama/format farklarını giderin (%, ondalık).")

    if "süre farkı" in dn or "duration" in kinds:
        # süre farkında birim bilgisini diff_note içinden çekmeye çalış
        unit = ""
        m = re.search(r"süre farkı\s*\((.*?)\)", dn)
        if m: unit = m.group(1)
        return ("E103",
                f"Süre değerleri{(' ('+unit+')') if unit else ''} uyuşmuyor" + (f" (örn: {sample})" if sample else ""),
                "Süre birimlerini standardize edin (gün/ay/yıl) ve mevzuattaki değere eşitleyin.")

    if "tarih farkı" in dn or "date" in kinds:
        return ("E104",
                "Tarih(ler) farklı",
                "PDF’deki tarih(leri) mevzuattaki geçerlilik/başlangıç-bitiş tarihleriyle uyumlayın.")

    # ---- 2) Kural yönü çelişkisi (izin/uygulanır ↔ yasak/uygulanmaz)
    if "kural yönü farkı" in dn or "polarity" in kinds or status == "conflict":
        return ("E106",
                "Kural yönü çelişkisi (izin/uygulanır ↔ yasak/uygulanmaz)",
                "Mevzuat hükmünü esas alın; PDF’yi ‘yasak/uygulanmaz/istisna’ yönüne göre yeniden yazın.")

    # ---- 3) Madde/bent atfı farklı
    # snippet'larda yakalanan madde kodları birbirinden farklıysa
    if madde_pdf != "-" or madde_law != "-":
        if madde_pdf != madde_law:
            return ("E105",
                    f"Atıf yapılan madde/bent farklı (PDF: {madde_pdf} ↔ Mevzuat: {madde_law})",
                    "PDF’de madde/bent referansını mevzuattaki doğru maddeye güncelleyin ya da metni doğru maddeye taşıyın.")

    # ---- 4) Benzerlik/eksiklik kaynaklı durumlar
    sim = float(row.get("similarity") or 0)
    if status == "low-similarity" or sim < 0.70:
        return ("E108",
                "Anlamsal benzerlik düşük; eşleşme zayıf",
                "Metni daha net/bağlamsal hâle getirin (başlık/madde kodu ekleyin); gerekirse chunk boyutlarını ve eşikleri yeniden ayarlayın.")

    if status == "missing":
        return ("E109",
                "Mevzuata karşılık gelen içerik PDF’de eksik",
                "Mevzuattaki ilgili hükmü PDF’ye ekleyin ya da referans verin.")

    if status == "ambiguous":
        return ("E110",
                "Eşleşme belirsiz (ambiguous)",
                "İnsan doğrulaması yapın; anahtar kelime/madde referanslarını güçlendirip tekrar kontrol edin.")

    if status == "weak_aligned":
        return ("E111",
                "Kısmi uyum (weak_aligned)",
                "Metni mevzuat terminolojisiyle güçlendirerek tam uyuma yaklaştırın.")

    # ---- 5) Varsayılan (yakalanamayan durumlar)
    return ("E000",
            "Fark türü belirgin değil",
            "İnsan doğrulaması yapın; diff_note ve eşleşen snippet’leri kontrol edin.")


# =============================================================================
# HTML rapor (pandas Styler ile)
# =============================================================================

def _style_row(s: pd.Series) -> List[str]:
    """
    HTML tablo satırlarını görsel olarak renklendirir (uyarı/önem/benzerlik/güven).
    - conflict → açık kırmızı
    - kritik   → daha koyu kırmızı
    - benzerlik < 0.88 → sarı
    - güven < 0.90     → gri
    """
    bg = ""
    try:
        if str(s.get("Durum", "")).lower() == "conflict":
            bg = "background-color:#ffe5e5;"
        if str(s.get("Önem", "")).lower() == "kritik":
            bg = "background-color:#ffd6d6;"
        sim = float(s.get("Benzerlik", 1))
        if sim < 0.88:
            bg = "background-color:#fff6cc;"
        conf = float(s.get("Güven", 1))
        if conf < 0.90:
            bg = "background-color:#f0f0f0;"
    except Exception:
        pass
    return [bg] * len(s)


def build_html_report(details_std: pd.DataFrame, overview: pd.DataFrame, law_version: str | None = None) -> str:
    """
    - Üstte 'En Sorunlu Sayfalar' tablosu
    - Altta sayfa sayfa problem listesi
    - Her satırda Kod/Tespit/Öneri gösterimi
    - Uyuşmayan değer çiftleri için collapsible mini tablo
    - (Yeni) Üstte mevzuat versiyon etiketi (varsa)
    """
    parts: List[str] = []
    parts.append("<h1>PDF Karşılaştırma Raporu</h1>")
    if law_version:
        parts.append(f"<p><b>Mevzuat versiyonu:</b> {law_version}</p>")

    # (A) Overview tablosu
    if not overview.empty:
        ov = overview.copy()
        ov = ov.rename(
            columns={
                "page": "Sayfa",
                "problems": "Sorun",
                "aligned": "Toplam Aligned",
                "total": "Toplam",
                "avg_sim": "Ort. Benzerlik",
                "problem_ratio": "Problem Oranı",
            }
        )
        parts.append("<h2>En Sorunlu Sayfalar</h2>")
        parts.append(
            ov.style
              .format({"Ort. Benzerlik": "{:.3f}", "Problem Oranı": "{:.3f}"})
              .hide(axis="index")
              .to_html()
        )

    # (B) Sayfa bazlı detaylar
    if not details_std.empty:
        det = details_std.rename(
            columns={
                "page": "Sayfa",
                "chunk_id": "Chunk",
                "status": "Durum",
                "severity": "Önem",
                "similarity": "Benzerlik",
                "confidence": "Güven",
                "diff_note": "Fark Notu",
                "explain_code": "Kod",
                "explain_text": "Tespit",
                "action_hint": "Öneri",
                "pdf_snippet": "PDF",
                "law_snippet": "Mevzuat",
                "pdf_idx": "PDF_İndeks",
                "law_idx": "Mevzuat_İndeks",
                "madde_pdf": "Madde(PDF)",
                "madde_law": "Madde(Mevzuat)",
                "ecode_name": "Kod Adı",
                "ecode_desc": "Kod Açıklaması",
            }
        ).copy()

        parts.append("<h2>Sayfa Bazlı Örnek Hatalar</h2>")
        for sayfa, grp in det.sort_values(["Sayfa"]).groupby("Sayfa", dropna=False):
            parts.append(f"<h3>Sayfa {sayfa}</h3>")
            cols = [c for c in [
                "Chunk","PDF_İndeks","Mevzuat_İndeks","Durum","Önem","Benzerlik","Güven",
                "Madde(PDF)","Madde(Mevzuat)","Kod","Kod Adı","Kod Açıklaması",
                "Tespit","Öneri","Fark Notu","PDF","Mevzuat"
            ] if c in grp.columns]

            parts.append(
                grp[cols]
                  .style.apply(_style_row, axis=1)
                  .hide(axis="index")
                  .set_properties(**{"white-space": "nowrap", "font-size": "12px"})
                  .to_html()
            )

            # Uyuşmayan değer çiftlerini küçük bir tabloyla sunalım (gerektiğinde açılır)
            for _, row in grp.iterrows():
                pdf_list = row.get("pdf_list")
                law_list = row.get("law_list")
                pairs = mismatch_pairs(pdf_list, law_list)
                if pairs:
                    sub = pd.DataFrame([{"PDF Değeri": a, "Mevzuat Değeri": b} for a, b in pairs])
                    parts.append(f'<details><summary>Chunk {row["Chunk"]} — Uyumsuz değerler</summary>')
                    parts.append(sub.style.hide(axis="index").to_html())
                    parts.append("</details>")

    return "\n".join(parts)


def save_html_report(details_std: pd.DataFrame, overview: pd.DataFrame, out_dir: Path, ts: str, law_version: str | None = None) -> Path:
    """HTML raporu diske yazar ve yolunu döndürür."""
    out_dir.mkdir(parents=True, exist_ok=True)
    html_str = build_html_report(details_std, overview, law_version=law_version)
    out = out_dir / f"analysis_{ts}.html"
    out.write_text(html_str, encoding="utf-8")
    return out


# =============================================================================
# (Opsiyonel) DOCX rapor — Office kullanıcıları için
# =============================================================================

def save_docx_report(details_std: pd.DataFrame, overview: pd.DataFrame, out_dir: Path, ts: str) -> Optional[Path]:
    """
    python-docx opsiyonel bağımlılık. Kurulu değilse sessizce None döndürür.
    Kuruluysa özet + sayfa detaylarını kısa formatta .docx’e döker.
    """
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    doc.add_heading('PDF Karşılaştırma Raporu', level=0)

    if not overview.empty:
        doc.add_heading('En Sorunlu Sayfalar', level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Sayfa"; hdr[1].text = "Sorun"; hdr[2].text = "Aligned"; hdr[3].text = "Toplam"
        hdr[4].text = "Problem Oranı"; hdr[5].text = "Ort. Benzerlik"
        for _, r in overview.iterrows():
            row = table.add_row().cells
            row[0].text = str(int(r["page"]))
            row[1].text = str(int(r["problems"]))
            row[2].text = str(int(r["aligned"]))
            row[3].text = str(int(r["total"]))
            row[4].text = f'{float(r["problems"]/r["total"]):.3f}'
            row[5].text = f'{0 if pd.isna(r["avg_sim"]) else float(r["avg_sim"]):.3f}'

    if not details_std.empty:
        doc.add_heading('Sayfa Bazlı Örnek Hatalar', level=1)
        for page, grp in details_std.groupby("page"):
            doc.add_heading(f"Sayfa {int(page)}", level=2)
            for _, r in grp.iterrows():
                p = doc.add_paragraph()
                p.add_run(
                    f'Chunk {r["chunk_id"]} • {r["status"]} • sim={r.get("similarity","")}, conf={r.get("confidence","")}'
                ).bold = True
                # Kod + Tespit + Öneri
                doc.add_paragraph(f'[{r.get("explain_code","")}] {r.get("explain_text","")}')
                doc.add_paragraph(f'Öneri: {r.get("action_hint","")}')
                # Uyumsuz değerler tablosu
                pairs = mismatch_pairs(r.get("pdf_list"), r.get("law_list"))
                if pairs:
                    t = doc.add_table(rows=1, cols=2)
                    t.rows[0].cells[0].text = "PDF Değeri"; t.rows[0].cells[1].text = "Mevzuat Değeri"
                    for a, b in pairs:
                        row = t.add_row().cells
                        row[0].text = str(a); row[1].text = str(b)

    path = out_dir / f"analysis_{ts}.docx"
    doc.save(str(path))
    return path


# =============================================================================
# KISA RAPOR (Yönetici özeti: yalnız Kod + Tespit + Öneri)
# =============================================================================

def resolve_law_version(cli_value: Optional[str]) -> str:
    """
    Yürürlük/versiyon etiketini belirler.
    Öncelik: --law_version > ENV(LAW_VERSION) > 'bilinmiyor'
    """
    if cli_value and str(cli_value).strip():
        return str(cli_value).strip()
    env = os.environ.get("LAW_VERSION")
    return env.strip() if env else "bilinmiyor"


def _ecode_meta(code: str) -> Tuple[str, str]:
    """E-kodu için (ad, açıklama) döndür. Bulunamazsa ('—','—')."""
    d = E_CODE_DICT.get((code or "").upper().strip())
    return (d["ad"], d["aciklama"]) if d else ("—", "—")


def save_short_report(details_std: pd.DataFrame, overview: pd.DataFrame, out_dir: Path, ts: str, law_version: str) -> Path:
    """
    Çok kısa, yönetime uygun rapor:
      - Üstte tarih ve mevzuat versiyonu
      - Her sayfa için yalnız: [E-KODU] Tespit — Öneri (tek satırlar)
      - (yeni) en altta kodlara göre dağılım özeti
    """
    lines: List[str] = []
    lines.append(f"# Kısa Rapor — {ts}")
    lines.append(f"- Mevzuat versiyonu: **{law_version}**")
    lines.append(f"- Toplam sorun satırı: **{int(overview['problems'].sum())}**\n")

    # Sayfa sırası: en sorunlu → az sorunlu
    page_order = overview.sort_values(
        ["problems", "problem_ratio", "page"],
        ascending=[False, False, True]
    )["page"].tolist()

    for p in page_order:
        sub = details_std[details_std["page"] == p]
        if sub.empty:
            continue
        lines.append(f"## Sayfa {int(p)}")
        for _, r in sub.iterrows():
            code = str(r.get("explain_code","") or "").strip()
            tesp = str(r.get("explain_text","") or r.get("explanation","") or "").strip()
            oneri= str(r.get("action_hint","") or "").strip()
            if not (code or tesp or oneri):
                continue
            lines.append(f"- **[{code or '—'}]** {tesp or '—'} — _Öneri:_ {oneri or '—'}")
        lines.append("")  # boş satır ayırıcı

    # (yeni) kod dağılımı
    counts = details_std["explain_code"].fillna("E000").value_counts().sort_index()
    if not counts.empty:
        lines.append("## Kodlara göre dağılım")
        lines.append("| Kod | Ad | Açıklama | Adet |")
        lines.append("|--|--|--|--:|")
        for code, n in counts.items():
            ad, acik = _ecode_meta(str(code))
            lines.append(f"| {code} | {ad} | {acik} | {int(n)} |")

    out = Path(out_dir) / f"analysis_{ts}_short.md"
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


# =============================================================================
# ANA ÇALIŞMA
# =============================================================================

def main():
    # (1) Argümanları topla
    ap = argparse.ArgumentParser()
    ap.add_argument("--csv", required=True, help="pdf_gap_checker çıktısı CSV yolu")
    ap.add_argument("--out", default="reports", help="çıktı klasörü (varsayılan: reports)")
    ap.add_argument(
        "--low_sim_threshold", type=float, default=0.70,
        help="Benzerlik bu eşikten düşükse 'low-similarity' olarak işaretlenir (default: 0.70)"
    )
    # (yeni) versiyon etiketi için opsiyonel parametre
    ap.add_argument("--law_version", default=None, help="Mevzuat versiyonu/yürürlük bilgisi (opsiyonel)")
    args = ap.parse_args()

    # (2) Zaman etiketi ve çıktı klasörü
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    os.makedirs(args.out, exist_ok=True)

    # (3) CSV oku
    df = pd.read_csv(args.csv, encoding="utf-8", low_memory=False)

    # (4) Kolonları esnekçe bul
    col_status   = smart_col(df, ["status"])
    col_severity = smart_col(df, ["severity"])
    col_conf     = smart_col(df, ["confidence"])
    col_sim      = smart_col(df, ["similarity"])
    col_page     = smart_col(df, ["pdf_page","page"])
    col_chunk_id = smart_col(df, ["pdf_chunk_id","chunk_id"])
    col_pdf_snip = smart_col(df, ["pdf_snippet","snippet_pdf"])
    col_law_snip = smart_col(df, ["law_snippet","snippet_law"])

    # (5) Opsiyonel kolonlar (varsa kullanılır)
    col_diff     = smart_col_opt(df, ["diff_note","diff","note"])
    col_pdf_idx  = smart_col_opt(df, ["pdf_idx","pdf_index","index_pdf"])
    col_law_idx  = smart_col_opt(df, ["law_idx","law_index","index_law"])

    # (6) Status normalizasyonu + düşük benzerliği otomatik etiketle
    df["_status_norm"] = df[col_status].astype(str).str.lower().str.strip()
    df["_similarity"]  = pd.to_numeric(df[col_sim], errors="coerce")
    df.loc[df["_similarity"] < args.low_sim_threshold, "_status_norm"] = "low-similarity"

    # (7) Problem/OK bayrakları
    df["_is_problem"] = df["_status_norm"].isin(PROBLEM_STATUSES)
    df["_is_ok"]      = df["_status_norm"].isin(OK_STATUSES)

    # (8) Özet (sayfa bazında grup — problems / aligned / total / avg_sim)
    grp = (
        df.groupby(df[col_page])
          .agg(
              problems=("_is_problem", "sum"),
              aligned=("_is_ok", "sum"),
              total=("_status_norm", "count"),
              avg_sim=("_similarity", "mean"),
          )
          .reset_index()
          .rename(columns={col_page: "page"})
    )
    grp["problem_ratio"] = (grp["problems"] / grp["total"]).round(3)
    grp = grp.sort_values(["problems", "problem_ratio", "page"], ascending=[False, False, True])

    # (9) Detay tablosu — sadece problem satırları
    details = df[df["_is_problem"]].copy()
    details["status"] = details["_status_norm"]

    # (10) İçe alacağımız kolonlar (varsa opsiyonelleri de ekle)
    keep_cols = [col_page, col_chunk_id, "status", col_severity, "_similarity", col_conf]
    if col_pdf_idx: keep_cols.append(col_pdf_idx)
    if col_law_idx: keep_cols.append(col_law_idx)
    if col_diff:    keep_cols.append(col_diff)
    keep_cols += [col_pdf_snip, col_law_snip]
    details = details[keep_cols].copy()

    # (11) Standardize isimler — raporun geri kalanında bu adlarla ilerleyeceğiz
    new_cols = ["page","chunk_id","status","severity","similarity","confidence"]
    if col_pdf_idx: new_cols.append("pdf_idx")
    if col_law_idx: new_cols.append("law_idx")
    if col_diff:    new_cols.append("diff_note")
    new_cols += ["pdf_snippet","law_snippet"]
    details.columns = new_cols

    # (12) Tip/biçim düzeltmeleri
    details["chunk_id"]   = details["chunk_id"].astype(str)
    details["severity"]   = details["severity"].astype(str)
    details["similarity"] = pd.to_numeric(details["similarity"], errors="coerce")
    details["confidence"] = pd.to_numeric(details["confidence"], errors="coerce")
    if "diff_note" in details.columns:
        details["diff_note"] = details["diff_note"].astype(str)

    details["pdf_snippet"] = details["pdf_snippet"].map(lambda s: trim(s, 240))
    details["law_snippet"] = details["law_snippet"].map(lambda s: trim(s, 240))

    # (13) Ek sütunlar:
    #  - snippet'lardan madde numarası çıkar
    details["madde_pdf"] = details["pdf_snippet"].apply(extract_madde)
    details["madde_law"] = details["law_snippet"].apply(extract_madde)

    #  - diff_note'tan listeleri çıkar (pdf_list, law_list)
    details[["pdf_list","law_list"]] = details.get("diff_note", pd.Series([""]*len(details))).apply(
        lambda dn: pd.Series(parse_diff_lists(dn))
    )

    #  - Son kullanıcıya yönelik Kod+Tespit+Öneri (3'lü) üret
    exps = details.apply(make_explanation, axis=1)
    details["explain_code"] = [e[0] for e in exps]
    details["explain_text"] = [e[1] for e in exps]
    details["action_hint"]  = [e[2] for e in exps]

    #  - (yeni) heuristic etiketleri (raporlamada arama/filtresi için faydalı)
    details["detected_kinds"] = details.apply(lambda r: ",".join(sorted(_kind_from_row(r))), axis=1)

    #  - (yeni) kod adı/açıklaması sütunları
    details["ecode_name"] = details["explain_code"].map(lambda c: _ecode_meta(c)[0])
    details["ecode_desc"] = details["explain_code"].map(lambda c: _ecode_meta(c)[1])

    #  - (geri uyumluluk) eski alan adı 'explanation' bekleyen UI’lar için kopya
    details["explanation"] = details["explain_text"]

    # =============================================================================
    # ÇIKTILAR
    # =============================================================================

    # (A) Markdown — hızlı paylaşım ve sürüm kontrol dostu
    law_version = resolve_law_version(args.law_version)  # (yeni) versiyon etiketi
    md_lines: List[str] = []
    md_lines.append(f"# PDF Karşılaştırma Özeti ({ts})\n")
    md_lines.append(f"- Kaynak CSV: `{os.path.basename(args.csv)}`")
    md_lines.append(f"- Mevzuat versiyonu: **{law_version}**")
    md_lines.append(f"- Sorun sayısı: **{int(grp['problems'].sum())}**  |  Toplam satır: **{int(grp['total'].sum())}**\n")
    md_lines.append("## En Sorunlu Sayfalar\n")
    md_lines.append("| Sayfa | Sorun | Aligned | Toplam | Problem Oranı | Ortalama Benzerlik |")
    md_lines.append("|--:|--:|--:|--:|--:|--:|")
    for _, r in grp.head(20).iterrows():
        avg_sim = 0 if pd.isna(r["avg_sim"]) else float(r["avg_sim"])
        md_lines.append(
            f"| {int(r['page'])} | {int(r['problems'])} | {int(r['aligned'])} | {int(r['total'])} | {r['problem_ratio']:.2f} | {avg_sim:.3f} |"
        )

    # (yeni) Kodlara göre sayım — hızlı yönetici özeti
    counts = details["explain_code"].fillna("E000").value_counts().sort_index()
    if not counts.empty:
        md_lines.append("\n## Kodlara göre dağılım\n")
        md_lines.append("| Kod | Ad | Açıklama | Adet |")
        md_lines.append("|--|--|--|--:|")
        for code, n in counts.items():
            ad, acik = _ecode_meta(str(code))
            md_lines.append(f"| {code} | {ad} | {acik} | {int(n)} |")

    md_lines.append("\n## Sayfa Sayfa Sorunlar (ilk 10 satırdan örnekler)\n")
    for p in grp["page"].tolist():
        sub_all = details[details["page"] == p]
        sub = sub_all.head(10)
        if sub.empty:
            continue
        md_lines.append(f"\n### Sayfa {int(p)} — {len(sub_all)} sorun")
        md_lines.append("| chunk | status | sev. | sim | conf | code | explanation | action_hint | diff_note | pdf_snippet | law_snippet |")
        md_lines.append("|--:|--|--|--:|--:|--|--|--|--|--|--|")
        for _, r in sub.iterrows():
            sev  = "" if pd.isna(r["severity"]) else str(r["severity"])
            sim  = "" if pd.isna(r["similarity"]) else f"{float(r['similarity']):.3f}"
            conf = "" if pd.isna(r["confidence"]) else f"{float(r['confidence']):.3f}"
            diff = "" if "diff_note" not in r.index or pd.isna(r["diff_note"]) else str(r["diff_note"])
            code = md_escape(str(r.get("explain_code","")))
            exp  = md_escape(str(r.get("explain_text","")))
            act  = md_escape(str(r.get("action_hint","")))
            pdf_snip = md_escape(str(r["pdf_snippet"]))
            law_snip = md_escape(str(r["law_snippet"]))
            chunk = "" if pd.isna(r["chunk_id"]) else str(r["chunk_id"])
            md_lines.append(
                f"| {chunk} | {r['status']} | {sev} | {sim} | {conf} | {code} | {exp} | {act} | {diff} | {pdf_snip} | {law_snip} |"
            )

    md_path = os.path.join(args.out, f"analysis_{ts}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    # (B) Excel — iki sheet (Overview + ProblemDetails)
    xlsx_path = os.path.join(args.out, f"analysis_{ts}.xlsx")
    with pd.ExcelWriter(xlsx_path, engine="xlsxwriter") as xw:
        # Overview
        grp.to_excel(xw, sheet_name="Overview", index=False)

        # ProblemDetails — yeni sütunları dahil edelim
        cols = [
            "page","chunk_id","status","severity","similarity","confidence",
            "explain_code","explain_text","action_hint",
            "ecode_name","ecode_desc",               # (yeni) kod meta
            "detected_kinds",                        # (yeni) sinyal etiketi
            "diff_note","pdf_snippet","law_snippet",
            "pdf_idx","law_idx","madde_pdf","madde_law"
        ]
        cols = [c for c in cols if c in details.columns]  # yoksa düşür
        details[cols].to_excel(xw, sheet_name="ProblemDetails", index=False)

        # QoL biçimleme
        for sheet in ["Overview","ProblemDetails"]:
            ws = xw.sheets[sheet]
            ws.freeze_panes(1, 0)  # başlığı sabitle
            ws.autofilter(0, 0, ws.dim_rowmax, ws.dim_colmax)  # filtreler

    # (C) HTML — Türkçe başlık + renk kodu + uyumsuz değer tabloları (+ versiyon etiketi)
    html_path = save_html_report(details.copy(), grp.copy(), Path(args.out), ts, law_version=law_version)

    # (D) DOCX — opsiyonel (python-docx varsa)
    docx_path = save_docx_report(details.copy(), grp.copy(), Path(args.out), ts)

    # (E) KISA RAPOR — yalnız Kod + Tespit + Öneri (+ kod dağılımı)
    short_md_path = save_short_report(details.copy(), grp.copy(), Path(args.out), ts, law_version)

    # (F) Konsola kısa özet
    print(f"[OK] Özet (MD): {md_path}")
    print(f"[OK] Excel     : {xlsx_path}")
    print(f"[OK] HTML      : {html_path}")
    if docx_path:
        print(f"[OK] DOCX      : {docx_path}")
    print(f"[OK] Kısa Rapor: {short_md_path}")


# (14) Standart Python CLI girişi
if __name__ == "__main__":
    main()
